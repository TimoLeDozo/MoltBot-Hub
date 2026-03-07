import datetime
import re
import time
from pathlib import Path

import streamlit as st
from docx import Document

from app_config import (
    DATA_ENTREE_DIR,
    DATA_SORTIE_DIR,
    TEMPLATE_PATH,
    ensure_runtime_directories,
    get_setup_issues,
)
from lecteur_OCR import extraire_fiche_renseignement
from main import analyser_grille_ia

ensure_runtime_directories()


def sauvegarder_fichier_upload(fichier_upload, prefixe_fichier):
    extension = Path(fichier_upload.name).suffix.lower() or ".bin"
    chemin = DATA_ENTREE_DIR / f"{prefixe_fichier}{extension}"
    with open(chemin, "wb") as fichier_sortie:
        fichier_sortie.write(fichier_upload.getbuffer())
    return str(chemin)


def generer_word(chemin_modele, chemin_sortie, donnees):
    doc = Document(chemin_modele)
    for paragraphe in doc.paragraphs:
        for cle, valeur in donnees.items():
            balise = f"{{{{{cle}}}}}"
            if balise in paragraphe.text:
                paragraphe.text = paragraphe.text.replace(balise, str(valeur))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraphe in cell.paragraphs:
                    for cle, valeur in donnees.items():
                        balise = f"{{{{{cle}}}}}"
                        if balise in paragraphe.text:
                            paragraphe.text = paragraphe.text.replace(balise, str(valeur))
    doc.save(str(chemin_sortie))


st.set_page_config(page_title="Assistant Bilans Orthophoniques", page_icon="📝", layout="centered")
st.title("📝 Assistant Bilans Orthophoniques")
st.divider()

setup_issues = get_setup_issues()
if setup_issues:
    st.error("Configuration incomplete avant generation.")
    for issue in setup_issues:
        st.write(f"- {issue}")
    st.caption("Le script install_et_lance.bat cree le fichier .env au premier lancement si besoin.")

st.header("1. Documents du patient")
fichier_renseignement = st.file_uploader("Fiche de renseignement (PDF)", type=["pdf"])
col1, col2, col3 = st.columns(3)
with col1:
    num_secu = st.text_input("N° de Sécurité Sociale (si absent)")
with col2:
    trimestre = st.selectbox("Période", ["Non précisé", "1er trimestre", "2ème trimestre", "3ème trimestre"])
with col3:
    genre = st.radio("Sexe de l'enfant", ["Garçon", "Fille"])

st.header("2. Anamnèse")
motif = st.text_input("Motif de consultation", placeholder="Ex: Faire le point sur ses connaissances en CP.")
comportement = st.text_input("Comportement de l'enfant", placeholder="Ex: enfant coopérant")
plainte = st.text_area(
    "Plainte de la famille / école",
    placeholder="Ex: Pas de plainte de la famille, mais plainte de l'école concernant l'attention.",
)
col_a, col_b = st.columns(2)
with col_a:
    fratrie = st.text_input("Fratrie", placeholder="Ex: 2 enfants avec une soeur de 11 ans.")
    langue = st.text_input("Langues parlées à la maison", placeholder="Ex: Français et un peu de bambara.")
with col_b:
    interets = st.text_input("Centres d'intérêt")

scolarisation = st.text_area(
    "Historique et lieu de scolarisation",
    placeholder="Ex: Maternelle réalisée à Sarcelles. Scolarisation en CP à l'école Monge à Créteil.",
)
antecedents = st.text_area(
    "Antécédents médicaux (Vue, Ouïe, etc.)",
    placeholder="Ex: Vision et audition à tester.",
)

st.header("3. Synthèse & Diagnostic")
diag_global = st.text_area(
    "Diagnostic global",
    placeholder="Ex: DANTE présente des difficultés de mémorisation importantes de vocabulaire...",
)

col_synth_1, col_synth_2 = st.columns(2)
with col_synth_1:
    res_phono = st.text_area(
        "Résumé Phonologie / Mémoire",
        placeholder="Ex: Manque de mémoire de travail / difficultés à mémoriser...",
    )
    res_oral = st.text_area(
        "Résumé Langage Oral",
        placeholder="Ex: Syntaxe dans la moyenne. Vocabulaire entre -1 et -2 écarts.",
    )
with col_synth_2:
    res_ecrit = st.text_area(
        "Résumé Langage Écrit",
        placeholder="Ex: Le langage écrit est presque inexistant. Score en écriture à -3 écarts.",
    )
    res_maths = st.text_area(
        "Résumé Mathématiques",
        placeholder="Ex: La numération orale est possible jusqu'à 23. Pas d'additions réalisées.",
    )

orientations = st.text_area(
    "Orientations et Aménagements (PAP, ULIS...)",
    placeholder="Ex: DANTE aurait besoin d'une orientation en ULIS, pour apprendre à son rythme. En attendant, un PAP peut être demandé...",
)
frequence = st.text_input("Fréquence rééducation", placeholder="Ex: 2 séances par semaine")
objectifs = st.text_area(
    "Objectifs thérapeutiques / Supports de travail",
    placeholder="Ex: Supports de lettres et chiffres représentés sur son bureau à repasser. Syllabes à lire...",
)

st.divider()

st.header("4. Grilles de résultats")
type_bilan = st.selectbox(
    "Quel type d'examens a été passé ?",
    ["Langage oral", "Langage oral et écrit", "Langage oral et logico-mathématique"],
)
grille_oral = st.file_uploader("Grille Oral (PDF/Image)", type=["jpg", "jpeg", "png", "pdf"])
grille_ecrit = None
if type_bilan in ["Langage oral et écrit", "Langage oral et logico-mathématique"]:
    grille_ecrit = st.file_uploader("Grille Écrit/Maths (Optionnel)", type=["jpg", "jpeg", "png", "pdf"])

if fichier_renseignement and grille_oral:
    if st.button(
        "✨ Générer le compte-rendu Word",
        type="primary",
        use_container_width=True,
        disabled=bool(setup_issues),
    ):
        with st.spinner("Analyse en cours... (Des pauses de sécurité de 25s sont actives pour l'IA) ⏳"):
            donnees_finales = {
                "NUM_SECU": num_secu if num_secu else "[À COMPLÉTER]",
                "TRIMESTRE": trimestre,
                "MOTIF_CONSULTATION": motif if motif else "[À COMPLÉTER]",
                "COMPORTEMENT_ENFANT": comportement if comportement else "[À COMPLÉTER]",
                "PLAINTE_FAMILLE": plainte if plainte else "[À COMPLÉTER]",
                "FRATRIE": fratrie if fratrie else "[À COMPLÉTER]",
                "LANGUE_MAISON": langue if langue else "[À COMPLÉTER]",
                "ANTECEDENTS_MEDICAUX": antecedents if antecedents else "[À COMPLÉTER]",
                "INTERETS_ENFANT": interets if interets else "[À COMPLÉTER]",
                "PARCOURS_SCOLAIRE": scolarisation if scolarisation else "[À COMPLÉTER]",
                "ORIENTATIONS_PAP_ULIS": orientations if orientations else "[À COMPLÉTER]",
                "DIAGNOSTIC_GLOBAL": diag_global if diag_global else "[À COMPLÉTER]",
                "RESUME_PHONOLOGIE": res_phono if res_phono else "[À COMPLÉTER]",
                "RESUME_LANGAGE_ORAL": res_oral if res_oral else "[À COMPLÉTER]",
                "RESUME_LANGAGE_ECRIT": res_ecrit if res_ecrit else "[À COMPLÉTER]",
                "RESUME_MATHEMATIQUES": res_maths if res_maths else "[À COMPLÉTER]",
                "FREQUENCE_REEDUCATION": frequence if frequence else "[À COMPLÉTER]",
                "OBJECTIFS_THERAPEUTIQUES": objectifs if objectifs else "[À COMPLÉTER]",
                "OBSERVATION_SPATIO_TEMPORELLE": "[À COMPLÉTER]",
                "LATERALITE": "[À COMPLÉTER]",
                "RESUME_GRAPHISME": "[À COMPLÉTER]",
                "RESUME_LECTURE_COMPREHENSION": "[À COMPLÉTER]",
                "ECART_LABYRINTHE": "[À COMPLÉTER]",
                "ECART_TOTAL_LECTURE": "[À COMPLÉTER]",
                "ECART_QUALITE_LECTURE": "[À COMPLÉTER]",
                "ECART_VITESSE_LECTURE": "[À COMPLÉTER]",
                "ECART_COMP_LECTURE_TEXTE": "[À COMPLÉTER]",
                "PHRASE_COMP_LECTURE_TEXTE": "[À COMPLÉTER]",
                "ECART_TOTAL_ORTHO": "[À COMPLÉTER]",
                "PHRASE_ORTHO": "[À COMPLÉTER]",
                "ECART_TOTAL_MATHS": "[À COMPLÉTER]",
                "ERREURS_COMP_LEXICALE ": "[À COMPLÉTER]",
            }

            if genre == "Garçon":
                donnees_finales.update({"PRONOM_SUJET": "il", "PRONOM_SUJET_MAJ": "Il", "ACCORD_E": ""})
            else:
                donnees_finales.update({"PRONOM_SUJET": "elle", "PRONOM_SUJET_MAJ": "Elle", "ACCORD_E": "e"})

            donnees_patient = extraire_fiche_renseignement(
                sauvegarder_fichier_upload(fichier_renseignement, "fiche_renseignement")
            )

            date_str = donnees_patient.get("DATE_NAISSANCE", "")
            match_date = re.search(r"(\d{2})[/-](\d{2})[/-](\d{4})", date_str)
            if match_date:
                jour, mois, annee = map(int, match_date.groups())
                try:
                    date_naiss = datetime.date(annee, mois, jour)
                    aujourdhui = datetime.date.today()
                    age = aujourdhui.year - date_naiss.year - (
                        (aujourdhui.month, aujourdhui.day) < (date_naiss.month, date_naiss.day)
                    )
                    donnees_patient["AGE_PATIENT"] = f"{age} ans"
                except Exception:
                    pass

            donnees_finales.update(donnees_patient)
            donnees_finales["ERREURS_COMP_LEXICALE\xa0"] = "[À COMPLÉTER]"

            time.sleep(25)

            res_oral_ia = analyser_grille_ia(
                sauvegarder_fichier_upload(grille_oral, "grille_orale"),
                type_test="general",
            )
            donnees_finales.update(res_oral_ia)

            if grille_ecrit:
                time.sleep(25)
                res_ecrit_ia = analyser_grille_ia(
                    sauvegarder_fichier_upload(grille_ecrit, "grille_ecrit"),
                    type_test="lecture",
                )
                donnees_finales.update(res_ecrit_ia)

            nom_brut = donnees_finales.get("NOM_PATIENT", "Patient")
            for caractere in ['\\', '/', ':', '*', '?', '"', '<', '>', '|', '\n', '\r']:
                nom_brut = nom_brut.replace(caractere, "")

            nom_patient = nom_brut.replace(" ", "_").strip("_") or "Patient"
            chemin_sortie = DATA_SORTIE_DIR / f"CR_{nom_patient}.docx"

            generer_word(TEMPLATE_PATH, chemin_sortie, donnees_finales)

            st.success("🎉 Succès ! Le compte-rendu a été généré.")
            with open(chemin_sortie, "rb") as fichier_genere:
                st.download_button("📥 Télécharger le Bilan Word", fichier_genere, f"CR_{nom_patient}.docx")
